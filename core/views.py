from django.shortcuts import redirect, render
from .models import Document
from .forms import DocumentForm

from django.contrib import messages

from django.db import connection

import docx2txt
import docx

from langdetect import detect, detect_langs, DetectorFactory
from iso639 import languages
import pytesseract
import cv2

from deep_translator import GoogleTranslator
from googletrans import Translator
from itertools import cycle

DetectorFactory.seed = 0

#Instance
con = connection.cursor()
translator = Translator()

def my_view(request):
    
    global con, translator

    message = 'Select file (.docx, .jpg)'
    # Handle file upload
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            newdoc = Document(docfile=request.FILES['docfile'])
            newdoc.save()

            # Redirect to the document list after POST
            return redirect('my-view')
        else:
            message = 'The form is not valid. Fix the following error:'
    else:
        form = DocumentForm()  #form

    # Load documents for the list page
    documents = Document.objects.all()
    
    field_name = 'docfile'
    obj = Document.objects.last()
    field_object = Document._meta.get_field(field_name)
    latest_file = 'documents/'+str(field_object.value_from_object(obj))
    print(type(latest_file), latest_file)
    
    if latest_file.endswith('.jpg'):
        text = ''
        data = obj.docfile
        pytesseract.pytesseract.tesseract_cmd=r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        img = cv2.imread(latest_file)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        # custom_config = r'-l eng+jpn --psm 6'
        custom_config = r'-l afr+amh+ara+asm+aze+bel+ben+bod+bos+bre+bul+cat+ceb+ces+chi_sim+chi_tra+chr+cos+cym+dan+dan_frak+dzo+ell+eng+enm+epo+equ+est+eus+fao+fas+fil+fin+fra+frk+frm+fry+gla+gke+glg+grc+guj+hat+heb+hin+hrv+hun+hye+iku+ind+isl+ita+ita+old+jav+jpn+kan+kat+kaz+khm+kir+kmr+kor+kur+lao+lat+lav+lit+itz+mal+mar+mkd+mlt+mon+mri+msa+mya+nep+nld+nor+oci+ori+osd+pan+pol+por+pus+que+ron+rus+san+sin+slk+slv+snd+spa+sqi+srp+sun+swa+swe+syr+tam+tat+tel+tgk+tgl+tha+tir+ton+tur+uig+ukr+urd+uzb+vie+yid+yor --psm 6'
        text_from_image = pytesseract.image_to_string(thresh, config=custom_config)
        text_from_image = str(text_from_image).splitlines()
        

        translated_text_from_image = []
        for to_translate in text_from_image:
            if to_translate != '':
                #Add destination in this line
                translated_text_from_image.append(translator.translate(to_translate, dest='en'))
            else:
                pass
        print(text_from_image) #prints the list with the text in image       
        

        Languages_detected_from_text_image = []
        try:
            for i in text_from_image:
                if i != '':
                    # print(detect_langs(i))
                    lang = detect(i)
                    print(lang, end=' ')

                    # for j in lang:
                    #     if j != 'cy':
                    #         lang = str(j)[0:2]
                    #     else:
                    #         pass

                    Languages_detected_from_text_image.append(languages.get(alpha2= lang).name)     
                else:
                    pass
            #print(Languages_detected_from_text)

            # for i in Languages_detected_from_text:
            #     print(i)
            
            # print(detect_langs(txt))
            # txt =txt.splitlines()
        except Exception as e:
            print(e)
        
        context = {'documents': documents, 'form': form, 'message': message, 'file': latest_file.replace('documents/',''), 'text_detected': text_from_image, 'text_translated': translated_text_from_image, 'Language_detected': Languages_detected_from_text_image}
        # cv2.imshow('image',thresh)
        # cv2.waitKey(0)
        
        # for i in txt:
        #     text += i
    elif latest_file.endswith('.doc') or latest_file.endswith('.docx'):
        DetectorFactory.seed = 0
        doc = docx.Document(latest_file)
    
        text_file = []
        for para in doc.paragraphs:
            text_file.append(para.text)
        text = '\n'.join(text_file)

        text_from_file = []
        text_from_file = text.splitlines()
        
        

        translated_text_from_file = []
        for to_translate in text_from_file:
            if to_translate != '':
                
                #Add destination in this line eg=en,jp
                translated_text_from_file.append(translator.translate(to_translate, dest='en'))
            else:
                pass
        print(text_from_file) #prints the list with the text in image  

        Languages_detected_from_text_file = []
        try:
            for i in text_from_file:
                if i != '':
                    if i != 'cy':
                        lang = detect(i)
                        print(lang, end=' ')
                        Languages_detected_from_text_file.append(languages.get(alpha2= lang).name)     
                else:
                    pass
        except Exception as e:
            print(e)
        
        context = {'documents': documents, 'form': form, 'message': message, 'file': latest_file.replace('documents/',''), 'text_detected': text_from_file, 'text_translated': translated_text_from_file, 'Language_detected': Languages_detected_from_text_file}
    else:
        messages.error(request, "Unsupportive File Type.")
        return render(request,"base.html")
        # for i in data_split:
        #     print(detect(i))
        #print(detect_langs(data))
    # with open(latest_file,'rb') as f:
    #     context = f.read().decode(encoding="utf-8")
    #     f.close()
    # latest_file = Document.objects.last()
    # text = docx2txt.process('documents')
    # path = 'D:\College\Projects\ImagetoText\documents'
    # files = os.listdir(path)
    # paths = [os.path.join(path, basename) for basename in files]
    # latest_file = max(paths, key=os.path.getctime)

    # list_of_files = glob.glob('D:\College\Projects\ImagetoText\documents') 
    # latest_file = max(list_of_files, key=os.path.getctime)
    # print(latest_file)

    # for filename in os.listdir(os.getcwd()):
    #     with open(os.path.join(os.getcwd(), filename), 'r') as f:
    #         text = f.read()
    # Render list page with the documents and the form

    # Color_List = ['primary','secondary','success','danger','info']

    # zip_list = zip(text_from_image, cycle(Color_List)) if len(text_from_image) > len(Color_List) else zip(cycle(text_from_image), Color_List)
    
        
 
        
    return render(request, 'base.html', context)
